# -*- coding: utf-8 -*-
"""
DOCX 통합 보고서 생성기
- Part 1: 최종보고서.md (이미지 인라인 포함)
- Part 2: 분析_수치_근거.md (상세 수치 부록)
- 이미지 미참조분은 이미지 갤러리로 자동 추가
출력: report/통합보고서.docx

※ Word → PowerPoint 복붙이 잘 안 되는 이유:
  - python-docx 표·커스텀 구분선(OOXML)은 PPT 클립보드와 호환 약함
  - 발표용은 build_pptx_report.py → report/최종보고서_발표.pptx 사용 권장
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
REPORT_DIR = ROOT / "report"
IMG_DIR = REPORT_DIR / "images"
OUTPUT = REPORT_DIR / "통합보고서.docx"

# 파일 탐색 (한글 인코딩 안전)
def find_md(name_fragment):
    for f in REPORT_DIR.iterdir():
        if name_fragment in f.name and f.suffix == ".md":
            return f
    return None

MD_MAIN = REPORT_DIR / "최종보고서.md"
MD_DATA = find_md("수치_근거") or find_md("수치근거")


# ── 헬퍼 ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_inline_runs(paragraph, text: str):
    """마크다운 인라인(볼드, 이탤릭, 코드)을 run으로 분해."""
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|~~[^~]+~~)"
    for part in re.split(pattern, text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("~~") and part.endswith("~~"):
            run = paragraph.add_run(part[2:-2])
            run.font.strike = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
        else:
            paragraph.add_run(part)


def parse_table_rows(raw_lines):
    rows = []
    for line in raw_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # 구분선 행(|---|---|) 제거
        if all(re.match(r"^[-:]+$", c) for c in cells if c):
            continue
        rows.append(cells)
    return rows


def build_word_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for ri, row_cells in enumerate(rows):
        for ci in range(ncols):
            cell_text = row_cells[ci] if ci < len(row_cells) else ""
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            add_inline_runs(para, cell_text)
            if ri == 0:  # 헤더 행
                for run in para.runs:
                    run.bold = True
                set_cell_bg(cell, "D6E4F0")
    doc.add_paragraph()


def embed_image(doc, img_path: Path, alt: str = ""):
    if not img_path.exists():
        doc.add_paragraph(f"[이미지 없음: {img_path.name}]")
        return
    try:
        # 이미지 너비 = 본문 너비(14cm)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(img_path), width=Cm(14))
        if alt:
            cap = doc.add_paragraph(alt)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()
    except Exception as e:
        doc.add_paragraph(f"[이미지 오류: {img_path.name} — {e}]")


def apply_font(doc):
    for style in doc.styles:
        try:
            font = style.font
            if font.name is None or "Calibri" in (font.name or ""):
                font.name = "맑은 고딕"
        except Exception:
            pass
    # Normal 스타일 명시 설정
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(10)


# ── 마크다운 파서 ──────────────────────────────────────────────────────────

def parse_md_to_doc(doc, md_path: Path, tracked_images: set):
    """마크다운 파일을 파싱해 doc에 추가. 삽입된 이미지를 tracked_images에 기록."""
    text = md_path.read_text(encoding="utf-8-sig")
    lines = text.splitlines()

    in_code = False
    code_buf = []
    table_buf = []
    i = 0

    def flush_table():
        nonlocal table_buf
        if table_buf:
            rows = parse_table_rows(table_buf)
            build_word_table(doc, rows)
            table_buf = []

    while i < len(lines):
        line = lines[i]

        # ── 코드 블록 ──────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            flush_table()
            if in_code:
                if code_buf:
                    para = doc.add_paragraph("\n".join(code_buf))
                    para.paragraph_format.left_indent = Cm(1)
                    for run in para.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(8.5)
                        run.font.color.rgb = RGBColor(0x20, 0x20, 0x80)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── 표 행이 아닌 경우 → 누적 표 flush ─────────────────────────────
        if not line.startswith("|"):
            flush_table()

        # ── 헤딩 ──────────────────────────────────────────────────────────
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # ── 수평선 ────────────────────────────────────────────────────────
        elif re.match(r"^---+$", line.strip()) or re.match(r"^\*\*\*+$", line.strip()):
            add_hr(doc)

        # ── 표 행 ─────────────────────────────────────────────────────────
        elif line.startswith("|"):
            table_buf.append(line)

        # ── 이미지 ────────────────────────────────────────────────────────
        elif line.strip().startswith("!["):
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
            if m:
                alt = m.group(1)
                rel = m.group(2).lstrip("./")
                img_path = md_path.parent / rel
                if not img_path.exists():
                    # images/ 직접 탐색
                    img_path = IMG_DIR / img_path.name
                embed_image(doc, img_path, alt)
                tracked_images.add(img_path.name)

        # ── 글머리 기호 ───────────────────────────────────────────────────
        elif re.match(r"^[-*]\s+", line):
            para = doc.add_paragraph(style="List Bullet")
            add_inline_runs(para, re.sub(r"^[-*]\s+", "", line))

        # ── 번호 목록 ─────────────────────────────────────────────────────
        elif re.match(r"^\d+\.\s+", line):
            para = doc.add_paragraph(style="List Number")
            add_inline_runs(para, re.sub(r"^\d+\.\s+", "", line))

        # ── 인용구 ────────────────────────────────────────────────────────
        elif line.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            add_inline_runs(para, line[2:])
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
                run.italic = True

        # ── 빈 줄 ─────────────────────────────────────────────────────────
        elif line.strip() == "":
            pass

        # ── 일반 단락 ─────────────────────────────────────────────────────
        else:
            para = doc.add_paragraph()
            add_inline_runs(para, line)

        i += 1

    flush_table()


# ── 메인 ──────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    apply_font(doc)

    # 페이지 설정 (A4)
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    tracked = set()

    # ── Part 1: 최종보고서 ────────────────────────────────────────────────
    if MD_MAIN.exists():
        print(f"[1/2] {MD_MAIN.name} 변환 중...")
        parse_md_to_doc(doc, MD_MAIN, tracked)
        doc.add_page_break()
    else:
        print(f"[경고] {MD_MAIN} 없음, 건너뜀")

    # ── Part 2: 분析_수치_근거 ────────────────────────────────────────────
    if MD_DATA and MD_DATA.exists():
        print(f"[2/2] {MD_DATA.name} 변환 중...")
        h = doc.add_heading("부록. 분석 수치 근거 상세", level=1)
        parse_md_to_doc(doc, MD_DATA, tracked)
        doc.add_page_break()
    else:
        print("[경고] 분析_수치_근거.md 없음, 건너뜀")

    # ── 이미지 갤러리 (미삽입분) ──────────────────────────────────────────
    all_imgs = sorted(IMG_DIR.glob("*.png"))
    missing = [p for p in all_imgs if p.name not in tracked]

    if missing:
        doc.add_heading("이미지 갤러리 (전체)", level=1)
        for img_path in missing:
            label = img_path.stem.replace("_", " ")
            doc.add_heading(label, level=3)
            embed_image(doc, img_path, label)

    doc.save(str(OUTPUT))
    print(f"\n[완료] 저장: {OUTPUT}")
    print(f"  이미지: {len(tracked)} 인라인 + {len(missing)} 갤러리")


if __name__ == "__main__":
    main()
